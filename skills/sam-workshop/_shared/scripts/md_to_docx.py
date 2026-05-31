#!/usr/bin/env python3
"""
md_to_docx.py — sam-workshop / G5 facilitator-run helper (v1.3 2026-05-03)

Markdown → DOCX regenerator (G5). Self-contained copy from the main pipeline's
`shared/scripts/md_to_docx.py` (v1.2 2026-05-03). Workshop pack ships
self-contained — code lives inside the workshop repo, not referenced
cross-pipeline.

Single-shot helper that combines:
  1. pandoc (Quarto-bundled binary or PATH) for md → docx conversion
  2. python-docx post-processing for TNR 12pt + 1in margins +
     optional double line spacing

Workshop integration (v1.3):
  - Step 8b2 wrap — **facilitator-only** call after desk-reject-precheck
    Pass/Fail decision is recorded. Participant does NOT call this directly
    (kept off the 17-skill mental model).
  - Output: `paper_home/08_package/manuscript_full.docx` automatically
    generated from `paper_home/04_draft/manuscript.md`.
  - On failure (pandoc missing, etc): facilitator notifies participant that
    docx packaging is left as homework. Workshop flow continues.

Usage:
    # Facilitator-run from workshop wrap step:
    python skills/sam-workshop/_shared/scripts/md_to_docx.py \\
        --md   paper_home/04_draft/manuscript.md \\
        --docx paper_home/08_package/manuscript_full.docx \\
        --double-spaced

Programmatic:
    from md_to_docx import regenerate_docx_from_md
    out = regenerate_docx_from_md(
        md_path="paper_home/04_draft/manuscript.md",
        docx_path="paper_home/08_package/manuscript_full.docx",
        double_line_spacing=True,
    )

Dependencies:
  - pandoc (Quarto bundle preferred on Windows). Install via
    https://quarto.org/docs/get-started/.
  - python-docx (`pip install python-docx`)
"""
from __future__ import annotations

import os
import subprocess
import sys
from pathlib import Path


# Pandoc auto-discovery: prefer Quarto bundle on Windows; fall back to PATH.
_PANDOC_CANDIDATES = [
    r"C:\Program Files\Quarto\bin\tools\pandoc.exe",
    r"C:\Program Files (x86)\Quarto\bin\tools\pandoc.exe",
    "/Applications/quarto/bin/tools/pandoc",
    "/usr/local/bin/pandoc",
    "/usr/bin/pandoc",
]


class FacilitatorActionRequired(RuntimeError):
    """Workshop-friendly exception: facilitator should switch to homework
    fallback rather than try to fix in front of participants."""


def find_pandoc() -> str:
    """Locate the pandoc binary. Returns absolute path or 'pandoc' if it's
    discoverable on PATH only. Raises FacilitatorActionRequired if no
    candidate is executable so the workshop can fall back gracefully."""
    env_override = os.environ.get("PANDOC")
    if env_override:
        if Path(env_override).is_file():
            return env_override
    for c in _PANDOC_CANDIDATES:
        if Path(c).is_file():
            return c
    # Last resort: trust PATH
    try:
        out = subprocess.run(["pandoc", "--version"],
                              capture_output=True, text=True, timeout=5)
        if out.returncode == 0:
            return "pandoc"
    except Exception:
        pass
    raise FacilitatorActionRequired(
        "pandoc not found. Set PANDOC env var to a pandoc binary path, "
        "or install Quarto (which bundles pandoc). "
        "Workshop fallback: leave .docx packaging as homework.")


def _run_pandoc(md_path: Path, docx_path: Path,
                 reference_doc: Path | None = None) -> None:
    cmd = [find_pandoc(),
           "--from=markdown",
           "--to=docx",
           "--output", str(docx_path),
           str(md_path)]
    if reference_doc:
        cmd.extend(["--reference-doc", str(reference_doc)])
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
    if proc.returncode != 0:
        raise FacilitatorActionRequired(
            f"pandoc failed (rc={proc.returncode}):\n"
            f"  stdout: {proc.stdout.strip()[:500]}\n"
            f"  stderr: {proc.stderr.strip()[:500]}\n"
            "Workshop fallback: leave .docx packaging as homework.")


def _apply_typography(docx_path: Path, *,
                       font: str = "Times New Roman",
                       size_pt: int = 12,
                       margins_in: float = 1.0,
                       double_line_spacing: bool = False) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_LINE_SPACING
    except ImportError:
        raise FacilitatorActionRequired(
            "python-docx not installed. Run `pip install python-docx`. "
            "Workshop fallback: leave typography polish as homework.")

    doc = Document(str(docx_path))

    for s in doc.sections:
        s.top_margin = Inches(margins_in)
        s.bottom_margin = Inches(margins_in)
        s.left_margin = Inches(margins_in)
        s.right_margin = Inches(margins_in)

    for para in doc.paragraphs:
        if double_line_spacing:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in para.runs:
            run.font.name = font
            run.font.size = Pt(size_pt)

    # Apply to table cells too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if double_line_spacing:
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    for run in para.runs:
                        run.font.name = font
                        run.font.size = Pt(size_pt)

    doc.save(str(docx_path))


def regenerate_docx_from_md(*,
                              md_path: Path | str,
                              docx_path: Path | str,
                              font: str = "Times New Roman",
                              size_pt: int = 12,
                              margins_in: float = 1.0,
                              double_line_spacing: bool = False,
                              reference_doc: Path | str | None = None) -> Path:
    """End-to-end: pandoc convert + typography post-process.

    Returns the absolute path to the generated .docx.

    Args:
        md_path: source Markdown file.
        docx_path: target .docx path.
        font: font family name (default Times New Roman).
        size_pt: font size in points (default 12).
        margins_in: page margins in inches (default 1.0, all four sides).
        double_line_spacing: apply WD_LINE_SPACING.DOUBLE to all paragraphs.
            Defaults to False because cover letters are typically single-spaced;
            manuscripts use True.
        reference_doc: optional pandoc --reference-doc for style baseline.

    Raises:
        FileNotFoundError if md_path doesn't exist.
        FacilitatorActionRequired if pandoc/python-docx unavailable —
            facilitator should switch to homework fallback.
    """
    md_p = Path(md_path)
    docx_p = Path(docx_path)
    if not md_p.exists():
        raise FileNotFoundError(f"markdown source not found: {md_p}")

    docx_p.parent.mkdir(parents=True, exist_ok=True)
    _run_pandoc(md_p, docx_p,
                reference_doc=Path(reference_doc) if reference_doc else None)
    _apply_typography(docx_p,
                       font=font,
                       size_pt=size_pt,
                       margins_in=margins_in,
                       double_line_spacing=double_line_spacing)
    return docx_p.resolve()


# Convenience wrappers for common artifacts in the workshop pipeline.

def regenerate_manuscript_full(paper_home: Path | str,
                                 *,
                                 manuscript_md: Path | str | None = None,
                                 font: str = "Times New Roman",
                                 size_pt: int = 12,
                                 margins_in: float = 1.0) -> Path:
    """Regen `paper_home/08_package/manuscript_full.docx` from
    `paper_home/04_draft/manuscript.md` (or explicit override).
    Manuscripts are double-spaced per most journal specs."""
    home = Path(paper_home)
    md_p = Path(manuscript_md) if manuscript_md else home / "04_draft" / "manuscript.md"
    return regenerate_docx_from_md(
        md_path=md_p,
        docx_path=home / "08_package" / "manuscript_full.docx",
        font=font, size_pt=size_pt, margins_in=margins_in,
        double_line_spacing=True,
    )


def regenerate_cover_letter(paper_home: Path | str,
                             *,
                             font: str = "Times New Roman",
                             size_pt: int = 12,
                             margins_in: float = 1.0) -> Path:
    """Regen cover_letter.docx from cover_letter.md.
    Cover letters are single-spaced (default)."""
    home = Path(paper_home)
    return regenerate_docx_from_md(
        md_path=home / "08_package" / "cover_letter.md",
        docx_path=home / "08_package" / "cover_letter.docx",
        font=font, size_pt=size_pt, margins_in=margins_in,
        double_line_spacing=False,
    )


# ============================================================
# Preflight (D-30 / D-7 환경 검증, v1.3.1 신규)
# ============================================================

def _korean_font_check() -> tuple[bool, str]:
    """Best-effort Korean font availability check via the platform-typical
    paths. False negatives possible (font dirs vary); used as a hint only."""
    if sys.platform == "win32":
        candidates = [
            r"C:\Windows\Fonts\malgun.ttf",
            r"C:\Windows\Fonts\NanumGothic.ttf",
        ]
    elif sys.platform == "darwin":
        candidates = [
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/Library/Fonts/AppleSDGothicNeo.ttc",
        ]
    else:
        candidates = [
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        ]
    for c in candidates:
        if Path(c).is_file():
            return True, c
    return False, "no Korean font detected at typical paths (informational; pandoc will substitute)"


def _smoke_korean_docx() -> tuple[bool, str]:
    """Generate a tiny Korean docx in a temp dir to verify pandoc + python-docx
    + path encoding all work end to end. Returns (ok, message)."""
    import tempfile
    sample = "# 워크숍 v1.3.1 환경 검증\n\n환자는 65세 남성으로 30일 사망률은 4.8%였다.\n"
    try:
        with tempfile.TemporaryDirectory() as td:
            md_path = Path(td) / "preflight.md"
            docx_path = Path(td) / "preflight.docx"
            md_path.write_text(sample, encoding="utf-8")
            regenerate_docx_from_md(md_path=md_path, docx_path=docx_path,
                                    double_line_spacing=True)
            if docx_path.exists() and docx_path.stat().st_size > 0:
                return True, f"sample docx {docx_path.stat().st_size} bytes"
            return False, "docx written but empty"
    except FacilitatorActionRequired as e:
        return False, str(e)
    except Exception as e:
        return False, f"unexpected: {e}"


def run_preflight() -> int:
    """v1.3.1 facilitator preflight — call this on the workshop laptop
    BEFORE D-30 dry-run AND BEFORE D-0. Returns 0 = all OK, 1 = blocking
    failure (workshop should announce homework fallback up-front), 2 =
    warnings only (workshop can proceed but flag specific risks).
    """
    failures: list[str] = []
    warnings: list[str] = []
    print("md_to_docx preflight (v1.3.1)")
    print("=" * 60)

    # Python version
    py_ok = sys.version_info >= (3, 9)
    print(f"  [{'OK' if py_ok else 'WARN'}] Python {sys.version.split()[0]}")
    if not py_ok:
        warnings.append("Python < 3.9 (recommend 3.11+)")

    # python-docx
    try:
        import docx  # noqa: F401
        print(f"  [OK] python-docx import")
    except ImportError:
        print(f"  [FAIL] python-docx not installed (pip install python-docx)")
        failures.append("python-docx not installed")

    # pandoc
    try:
        pandoc = find_pandoc()
        try:
            out = subprocess.run([pandoc, "--version"],
                                  capture_output=True, text=True, timeout=5)
            ver = out.stdout.strip().split("\n")[0] if out.returncode == 0 else "unknown"
        except Exception:
            ver = "unknown"
        print(f"  [OK] pandoc found: {pandoc}  ({ver})")
    except FacilitatorActionRequired as e:
        print(f"  [FAIL] pandoc: {e}")
        failures.append("pandoc not found")

    # Korean font (best effort)
    ok, msg = _korean_font_check()
    print(f"  [{'OK' if ok else 'WARN'}] Korean font: {msg}")
    if not ok:
        warnings.append("Korean font check inconclusive")

    # Write permission to workshop's typical output area
    try:
        import tempfile
        td = Path(tempfile.gettempdir()) / "md_to_docx_preflight_write_test"
        td.mkdir(exist_ok=True)
        test_file = td / "write_check.txt"
        test_file.write_text("ok", encoding="utf-8")
        test_file.unlink()
        td.rmdir()
        print(f"  [OK] write permission: {tempfile.gettempdir()}")
    except Exception as e:
        print(f"  [FAIL] write permission: {e}")
        failures.append("write permission denied")

    # End-to-end Korean docx smoke (only if no failures so far)
    if not failures:
        ok, msg = _smoke_korean_docx()
        print(f"  [{'OK' if ok else 'FAIL'}] Korean docx smoke: {msg}")
        if not ok:
            failures.append(f"end-to-end smoke failed: {msg}")

    print("=" * 60)
    if failures:
        print(f"BLOCKING FAILURES ({len(failures)}):")
        for f in failures:
            print(f"  - {f}")
        print()
        print("Workshop fallback: announce 'docx 패키징은 homework로 진행' at outset.")
        return 1
    if warnings:
        print(f"WARNINGS ({len(warnings)}):")
        for w in warnings:
            print(f"  - {w}")
        print()
        print("Workshop can proceed but facilitator should monitor for related issues.")
        return 2
    print("All checks passed. md_to_docx ready for facilitator-run at Step 8b2.")
    return 0


# ============================================================
# CLI
# ============================================================

def main() -> None:
    import argparse
    p = argparse.ArgumentParser(
        prog="md_to_docx",
        description="sam-workshop G5 facilitator-run pandoc + TNR12 docx regenerator")
    p.add_argument("--preflight", action="store_true",
                   help="v1.3.1: run environment health check (pandoc / python-docx / "
                        "Korean font / write permission / end-to-end smoke). "
                        "Exit 0 = ready, 1 = blocking, 2 = warnings only. "
                        "Run before D-30 dry-run and before D-0.")
    p.add_argument("--md", help="Markdown source path")
    p.add_argument("--docx", help="Target .docx path")
    p.add_argument("--paper-home",
                   help="Convenience: regen manuscript_full.docx for a paper_home dir "
                        "(uses 04_draft/manuscript.md → 08_package/manuscript_full.docx). "
                        "Mutually exclusive with --md/--docx.")
    p.add_argument("--font", default="Times New Roman")
    p.add_argument("--size-pt", type=int, default=12)
    p.add_argument("--margins-in", type=float, default=1.0)
    p.add_argument("--double-spaced", action="store_true",
                   help="apply double line spacing (manuscripts)")
    p.add_argument("--reference-doc", default=None)
    args = p.parse_args()

    if args.preflight:
        raise SystemExit(run_preflight())

    try:
        if args.paper_home:
            out = regenerate_manuscript_full(
                args.paper_home,
                font=args.font, size_pt=args.size_pt, margins_in=args.margins_in)
        else:
            if not (args.md and args.docx):
                p.error("either --paper-home OR (--md AND --docx) required")
            out = regenerate_docx_from_md(
                md_path=args.md,
                docx_path=args.docx,
                font=args.font,
                size_pt=args.size_pt,
                margins_in=args.margins_in,
                double_line_spacing=args.double_spaced,
                reference_doc=args.reference_doc,
            )
        print(f"OK  {out}")
    except FacilitatorActionRequired as e:
        # Workshop-friendly: clear actionable message, exit 2 (≠ 0/1)
        print(f"[FACILITATOR_FALLBACK] {e}", file=sys.stderr)
        raise SystemExit(2)


if __name__ == "__main__":
    main()
